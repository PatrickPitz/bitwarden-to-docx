import argparse
import json
import pathlib

from docx import Document


def create_arg_parser() -> argparse.ArgumentParser:
    """
    Creates an argument parser for command line arguments.

    Returns:
        argparse.ArgumentParser: The created argument parser.
    """
    new_parser = argparse.ArgumentParser(description="Create a PDF-File from a Bitwarden .json export")
    new_parser.add_argument("-i",
                            "--input",
                            help="Path to the Bitwarden .json export file",
                            required=True,
                            type=pathlib.Path)
    new_parser.add_argument("-o",
                            "--output",
                            help="Path to the output PDF file",
                            required=True,
                            type=pathlib.Path)
    new_parser.add_argument("-fc",
                            "--filter-collection",
                            help="Filter for a specific collectionId",
                            required=False,
                            type=str)
    new_parser.add_argument("-fo",
                            "--filter-organization",
                            help="Filter for a specific organizationId",
                            required=False,
                            type=str)
    new_parser.add_argument("-fon",
                            "--filter-name",
                            help="Filter for a organizations containing a specific string",
                            required=False,
                            type=str)
    new_parser.add_argument("-fu",
                            "--filter-username",
                            help="Filter items by username containing a specific string",
                            required=False,
                            type=str)
    new_parser.add_argument("-cor",
                            "--combined-or",
                            help="Filters are combined with OR; Default is AND",
                            required=False,
                            action="store_true")
    return new_parser


def extract_data_from_json(json_file: pathlib.Path) -> dict:
    with open(json_file, "r") as file:
        extracted_data = json.load(file)
    return extracted_data


def filter_items_and(items: list, filter_collection: str, filter_organization: str, filter_name: str,
                     filter_username: str) -> list:
    filtered_items = items
    if filter_collection:
        filtered_items = [item for item in filtered_items if filter_collection in item["collectionIds"]]
    if filter_organization:
        filtered_items = [item for item in filtered_items if item["organizationId"] == filter_organization]
    if filter_name:
        filtered_items = [item for item in filtered_items if filter_name in item["name"]]
    if filter_username:
        filtered_items = [item for item in filtered_items if
                          "login" in item and filter_username in item["login"]["username"]]
    return filtered_items


def filter_items_or(items: list, filter_collection: str, filter_organization: str, filter_name: str,
                    filter_username: str) -> list:
    filtered_items = []
    if filter_collection:
        filtered_items.extend([item for item in items if filter_collection in item["collectionIds"]])
    if filter_organization:
        filtered_items.extend([item for item in items if item["organizationId"] == filter_organization])
    if filter_name:
        filtered_items.extend([item for item in items if filter_name in item["name"]])
    if filter_username:
        filtered_items.extend(
            [item for item in items if "login" in item and filter_username in item["login"]["username"]])
    return filtered_items


if __name__ == '__main__':
    parser: argparse.ArgumentParser = create_arg_parser()
    args = parser.parse_args()
    items = extract_data_from_json(args.input)["items"]

    # Filtering items based on command line arguments

    if args.combined_or:
        items = filter_items_or(items, args.filter_collection, args.filter_organization, args.filter_name,
                                 args.filter_username)
    else:
        items = filter_items_and(items, args.filter_collection, args.filter_organization, args.filter_name,
                                args.filter_username)

    # Creating a Document and adding items to it
    document: Document = Document()
    document.add_heading("Bitwarden Export", level=1)
    for item in items:
        document.add_heading(item["name"], level=2)
        if "login" in item:
            document.add_paragraph(f"URL: {item['login']['uris'][0]['uri']}")
            document.add_paragraph(f"Username: {item['login']['username']}")
            document.add_paragraph(f"Password: {item['login']['password']}")
        document.add_paragraph(f"Notes: {item['notes']}")
    document.save(args.output)
